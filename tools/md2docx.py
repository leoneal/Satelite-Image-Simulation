#!/usr/bin/env python3
"""Convert the simulation report Markdown to a well-formatted Word document."""

import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── helpers ──────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_paragraph_border(paragraph, color="CCCCCC", sz="6"):
    """Add a left border to a paragraph (for blockquotes)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="{sz}" w:space="8" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val","single")}" '
            f'w:sz="{val.get("sz","4")}" w:space="0" w:color="{val.get("color","999999")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_code_run(paragraph, text, size=9):
    """Add a monospace run."""
    run = paragraph.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return run

def add_normal_run(paragraph, text, bold=False, italic=False, size=10.5):
    """Add a normal run with formatting."""
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return run

def parse_inline(text, paragraph, base_size=10.5):
    """Parse inline markdown: **bold**, `code`, *italic* and add runs."""
    # Order matters: code first (backticks), then bold, then italic
    pattern = re.compile(r'(`[^`]+`|\*\*.+?\*\*|\*.+?\*|\[.+?\]\(.+?\)|[^`*\[\n]+|\n)')
    parts = pattern.findall(text)
    for part in parts:
        if part == '\n':
            continue
        if part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(base_size)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.size = Pt(base_size)
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.font.size = Pt(base_size)
            run.italic = True
        elif part.startswith('[') and '](' in part:
            m = re.match(r'\[(.+?)\]\((.+?)\)', part)
            if m:
                run = paragraph.add_run(m.group(1))
                run.font.size = Pt(base_size)
                run.font.color.rgb = RGBColor(0x03, 0x69, 0xD6)
                run.underline = True
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(base_size)
    return paragraph

# ── main converter ───────────────────────────────────────────────────

def convert_md_to_docx(md_path, docx_path):
    doc = Document()

    # -- page setup --
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # -- styles --
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.35
    # Set East Asian font
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    for i in range(1, 5):
        h_style = doc.styles[f'Heading {i}']
        h_style.font.name = '微软雅黑'
        h_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        h_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        h_style.font.bold = True
        if i == 1:
            h_style.font.size = Pt(22)
            h_style.paragraph_format.space_before = Pt(24)
            h_style.paragraph_format.space_after = Pt(12)
        elif i == 2:
            h_style.font.size = Pt(16)
            h_style.paragraph_format.space_before = Pt(20)
            h_style.paragraph_format.space_after = Pt(8)
            # Add bottom border for H2
        elif i == 3:
            h_style.font.size = Pt(13)
            h_style.paragraph_format.space_before = Pt(16)
            h_style.paragraph_format.space_after = Pt(6)
        elif i == 4:
            h_style.font.size = Pt(11)
            h_style.paragraph_format.space_before = Pt(12)
            h_style.paragraph_format.space_after = Pt(4)

    # -- read file --
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Blank line
        if not stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Horizontal rule
        if stripped == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            # Add bottom border to simulate <hr>
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="12" w:space="1" w:color="CCCCCC"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue

        # Code block
        if stripped.startswith('```'):
            lang = stripped[3:].strip()
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i].rstrip('\n'))
                i += 1
            i += 1  # skip closing ```

            # Add a shaded code block paragraph
            for cl in code_lines:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.left_indent = Cm(0.5)
                # Grey background via shading
                pPr = p._p.get_or_add_pPr()
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="F0F0F0"/>')
                pPr.append(shd)
                run = p.add_run(cl if cl else ' ')
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            # Small gap after
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            continue

        # Table
        if stripped.startswith('|') and stripped.endswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1

            # Parse table
            rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                rows.append(cells)

            # Skip separator row (---|---)
            data_rows = [rows[0]]  # header
            for r in rows[1:]:
                if not all(re.match(r'^[-:]+$', c) for c in r):
                    data_rows.append(r)

            num_cols = len(data_rows[0])
            table = doc.add_table(rows=len(data_rows), cols=num_cols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for ri, row_data in enumerate(data_rows):
                for ci, cell_text in enumerate(row_data):
                    cell = table.cell(ri, ci)
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    if ri == 0:
                        # Header row
                        set_cell_shading(cell, "1A1A2E")
                        run = p.add_run(cell_text)
                        run.font.size = Pt(9.5)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.name = '微软雅黑'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    else:
                        set_cell_shading(cell, "FAFAFA" if ri % 2 == 0 else "FFFFFF")
                        parse_inline(cell_text, p, base_size=9.5)
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            continue

        # Standalone image: ![alt](path) — embed centered with caption
        img_match = re.match(r'^!\[(.+?)\]\((.+?)\)$', stripped)
        if img_match:
            alt = img_match.group(1)
            rel = img_match.group(2)
            img_path = os.path.join(os.path.dirname(md_path),
                                    rel.replace('/', os.sep))
            if os.path.exists(img_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, width=Cm(15.5))
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap.add_run(alt)
                cap_run.font.size = Pt(9)
                cap_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            else:
                print(f"WARNING: image not found: {img_path}")
            i += 1
            continue

        # Blockquote
        if stripped.startswith('>'):
            bq_lines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                bq_lines.append(lines[i].strip()[1:].strip())
                i += 1
            for bq in bq_lines:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                set_paragraph_border(p)
                parse_inline(bq, p, base_size=10)
                # Italic grey text for blockquotes
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                    run.italic = True
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            continue

        # Headings
        h1_match = re.match(r'^# (.+)', stripped)
        h2_match = re.match(r'^## (.+)', stripped)
        h3_match = re.match(r'^### (.+)', stripped)
        h4_match = re.match(r'^#### (.+)', stripped)

        if h1_match:
            p = doc.add_heading(h1_match.group(1), level=1)
            # Center the title
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        if h2_match:
            doc.add_heading(h2_match.group(1), level=2)
            i += 1
            continue
        if h3_match:
            doc.add_heading(h3_match.group(1), level=3)
            i += 1
            continue
        if h4_match:
            doc.add_heading(h4_match.group(1), level=4)
            i += 1
            continue

        # Unordered list
        if re.match(r'^- |^  - ', stripped):
            list_lines = []
            while i < len(lines) and re.match(r'^(  )?- ', lines[i].rstrip('\n')):
                list_lines.append(lines[i].rstrip('\n'))
                i += 1
            for ll in list_lines:
                m = re.match(r'^(  )?- (.+)', ll)
                if m:
                    indent_level = 1 if m.group(1) else 0
                    content = m.group(2)
                    p = doc.add_paragraph(style='List Bullet')
                    if indent_level > 0:
                        p.paragraph_format.left_indent = Cm(2)
                    p.clear()
                    parse_inline(content, p)
            continue

        # Ordered list
        if re.match(r'^\d+\. ', stripped):
            list_lines = []
            while i < len(lines) and re.match(r'^\d+\. ', lines[i].strip()):
                list_lines.append(lines[i].strip())
                i += 1
            for ll in list_lines:
                content = re.sub(r'^\d+\.\s*', '', ll)
                p = doc.add_paragraph(style='List Number')
                p.clear()
                parse_inline(content, p)
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        parse_inline(stripped, p)
        i += 1

    # -- save --
    doc.save(docx_path)
    print(f"Saved: {docx_path}")


if __name__ == '__main__':
    import sys
    if len(sys.argv) > 1:
        md_path = os.path.abspath(sys.argv[1])
    else:
        base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        md_path = os.path.join(base, 'docs', '仿真方案说明报告.md')
    docx_path = os.path.splitext(md_path)[0] + '.docx'
    convert_md_to_docx(md_path, docx_path)
